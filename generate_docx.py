import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# File paths for images
IMAGE_DIR = r"C:\Users\siddi\.gemini\antigravity\brain\4ea7a7cf-2b6a-4872-9ac6-bf2d86ef788f"
images = {
    "login": os.path.join(IMAGE_DIR, "login_screen_1781246946687.png"),
    "register": os.path.join(IMAGE_DIR, "register_screen_1781246961650.png"),
    "dashboard": os.path.join(IMAGE_DIR, "dashboard_screen_1781246975004.png"),
    "contacts": os.path.join(IMAGE_DIR, "contacts_screen_1781246989976.png"),
    "chat": os.path.join(IMAGE_DIR, "chat_screen_1781247002576.png"),
    "language": os.path.join(IMAGE_DIR, "language_screen_1781247016776.png"),
    "er_diagram": os.path.join(IMAGE_DIR, "er_diagram_1781247029233.png")
}

def create_report():
    doc = Document()
    
    # 1. Setup Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Set Default Normal Style (Times New Roman, 12pt, Double/1.15 line spacing)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Custom Heading Helper
    def add_custom_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_heading(level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0) # black
        
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        return heading

    # Custom Paragraph Helper
    def add_custom_paragraph(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = align
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = bold
            run.italic = italic
        return p

    # Custom Image Helper
    def add_custom_image(img_key, caption):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        
        img_path = images.get(img_key)
        if img_path and os.path.exists(img_path):
            try:
                run = p.add_run()
                run.add_picture(img_path, width=Inches(3.2))
            except Exception as e:
                print(f"Error loading image {img_key} ({img_path}): {e}")
                run = p.add_run(f"\n[Image Placeholder: {caption}]\n")
                run.font.name = 'Times New Roman'
                run.bold = True
        else:
            print(f"Image key '{img_key}' path does not exist: {img_path}")
            run = p.add_run(f"\n[Image Placeholder: {caption}]\n")
            run.font.name = 'Times New Roman'
            run.bold = True
            
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(12)
        
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.italic = True
        run_cap.font.color.rgb = RGBColor(80, 80, 80)

    # Custom Table Helper
    def add_custom_table(headers, rows_data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, h_text in enumerate(headers):
            hdr_cells[i].text = h_text
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                
        for row_data in rows_data:
            row = table.add_row()
            cells = row.cells
            for i, val in enumerate(row_data):
                cells[i].text = str(val)
                p = cells[i].paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                # Align short text / IDs / status to center
                if len(str(val)) < 12 or headers[i].lower() in ['tc id', 'status', 'user_id', 'contact_id', 'message_id', 'preferred_language', 'owner_id', 'sender_id', 'receiver_id']:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)

        # Space below table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(4)
        p_space.paragraph_format.space_after = Pt(4)

    # ==================== DOCUMENT CONTENT ====================
    
    # Title Header (Centered, bold)
    p_title = add_custom_paragraph("PROJECT REPORT ON SMART MESSENGER", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, bold=True)
    p_subtitle = add_custom_paragraph("REAL-TIME MULTILINGUAL MESSAGING SYSTEM", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, bold=True)
    
    # ------------------ ABSTRACT ------------------
    add_custom_heading("ABSTRACT", level=1, space_before=18)
    add_custom_paragraph(
        "Smart Messenger is a real-time multilingual messaging application designed to eliminate language barriers in digital communication. "
        "The application enables users to communicate with people speaking different languages by automatically translating messages into the "
        "recipient's preferred language. The system provides secure user authentication, contact management, instant messaging, language preference "
        "settings, and message translation services. Built using modern web and mobile technologies, the application ensures seamless communication "
        "across multiple languages while maintaining user privacy and message integrity. Smart Messenger offers an intuitive user interface and supports "
        "real-time conversations, making it useful for students, professionals, and users from diverse linguistic backgrounds. The application demonstrates "
        "the integration of messaging systems, cloud databases, authentication mechanisms, and translation APIs into a unified platform that enhances "
        "communication accessibility and user experience."
    )
    
    # Page break after abstract
    doc.add_page_break()
    
    # ------------------ CHAPTER 1 ------------------
    add_custom_heading("CHAPTER 1: INTRODUCTION", level=1)
    add_custom_heading("1.1 Smart Messenger – Real-Time Translator", level=2)
    add_custom_paragraph(
        "In today's globalized world, communication spans continents and languages. However, language barriers remain a significant hurdle in digital communication. "
        "The Smart Messenger application is a mobile messaging platform designed to solve this issue by allowing users to communicate across different languages using "
        "automatic, real-time machine translation technology. By combining a real-time chat application with robust language translation interfaces, it provides an "
        "accessible experience that translates incoming and outgoing messages instantly, facilitating natural conversations between users who do not share a common language."
    )
    
    add_custom_heading("1.2 Functionalities under User Login", level=2)
    add_custom_paragraph("The application provides a comprehensive suite of functionalities secured behind user login:")
    
    login_features = [
        "User Registration and Login: Allows new users to register and existing users to login to their profile.",
        "Secure Authentication: Integrates Firebase Authentication using email and password to secure user accounts.",
        "Profile Management: Allows users to configure their personal details including name, age, and profile picture.",
        "Contact Management: Enables search and management of chat contacts with specific language preferences.",
        "Real-Time Messaging: Supports instant message exchange powered by Cloud Firestore real-time sync.",
        "Automatic Message Translation: Automatically translates incoming messages into the recipient's preferred language.",
        "Language Preference Selection: Provides configuration options for both application UI language and preferred chat translation language.",
        "Chat History Storage: Safely persists and synchronizes chat logs and messages in the Firestore database.",
        "Speech Support: Integrates text-to-speech for speaking translated messages and speech-to-text for audio message dictation.",
        "User Dashboard: Serves as the central navigation hub showing active chats, profile options, and translation widgets."
    ]
    for idx, feature in enumerate(login_features, 1):
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.line_spacing = 1.15
        p_item.paragraph_format.space_after = Pt(3)
        run_num = p_item.add_run(f"{idx}. ")
        run_num.bold = True
        run_num.font.name = 'Times New Roman'
        run_text = p_item.add_run(feature)
        run_text.font.name = 'Times New Roman'
        
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

    # ------------------ CHAPTER 2 ------------------
    add_custom_heading("CHAPTER 2: HARDWARE AND SOFTWARE REQUIREMENTS", level=1)
    
    add_custom_heading("2.1 Hardware Requirements", level=2)
    hw_reqs = [
        "Processor: Intel Core i3 or above (Development) / Snapdragon 600 Series or above (Mobile device)",
        "RAM: Minimum 4 GB (8 GB recommended for development environment)",
        "Storage: Minimum 100 MB free space on device / 2 GB on development machine",
        "Network: Internet connection required for Firestore sync and translation services",
        "Android Device Version: Android 8.0 (Oreo) or above to run the compiled APK"
    ]
    for req in hw_reqs:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.line_spacing = 1.15
        p_item.paragraph_format.space_after = Pt(3)
        run = p_item.add_run(req)
        run.font.name = 'Times New Roman'
        
    add_custom_heading("2.2 Software Requirements", level=2)
    sw_reqs = [
        "Frontend Framework: React.js (built with Vite for fast build and loading speeds)",
        "Mobile Framework: Capacitor (to compile the web application into native Android wrapper)",
        "Backend Services: Google Firebase Console",
        "Authentication: Firebase Authentication (Email/Password, Google OAuth, and Phone Verification)",
        "Database: Google Cloud Firestore (noSQL real-time document store database)",
        "Translation API: Google Translate API (client-side single-query request endpoint)",
        "Integrated Development Environment (IDE): Visual Studio Code",
        "Mobile Build Tool: Android Studio (with Gradle for compilation and signing of APK)",
        "Programming Language: JavaScript (ES6+), HTML5, and CSS3"
    ]
    for req in sw_reqs:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.line_spacing = 1.15
        p_item.paragraph_format.space_after = Pt(3)
        run = p_item.add_run(req)
        run.font.name = 'Times New Roman'
        
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

    # ------------------ CHAPTER 3 ------------------
    add_custom_heading("CHAPTER 3: DESIGN LAYOUTS", level=1)
    add_custom_paragraph("This chapter outlines the key user interfaces of the Smart Messenger application, showing the layout design, stepper onboarding flow, and dashboard layout.")
    
    add_custom_heading("3.1 Login Screen", level=2)
    add_custom_paragraph("The Login Screen represents the entry point of the app (Step 1 of the onboarding wizard). It enables login via Email or Phone, and includes Google Sign-In.")
    add_custom_image("login", "Figure 3.1: Login Screen")
    
    add_custom_heading("3.2 Registration Screen", level=2)
    add_custom_paragraph("The Registration Screen (Step 4 of onboarding) collects the user's name and age to construct their database profile on the Firestore backend.")
    add_custom_image("register", "Figure 3.2: Registration Screen")
    
    add_custom_heading("3.3 Dashboard Screen", level=2)
    add_custom_paragraph("The Dashboard Screen is the main screen containing the search bar, active chat list, profile controls, and a quick 'Instant Translator' utility.")
    add_custom_image("dashboard", "Figure 3.3: Dashboard Screen")
    
    add_custom_heading("3.4 Contact List Screen", level=2)
    add_custom_paragraph("The Contact List displays saved chats and friends with their online status indicator and preferred translation language details.")
    add_custom_image("contacts", "Figure 3.4: Contact List Screen")
    
    add_custom_heading("3.5 Chat Screen", level=2)
    add_custom_paragraph("The Chat Screen provides real-time dialogue translation. Incoming messages are translated instantly and presented with original-translated text tabs and TTS audio icons.")
    add_custom_image("chat", "Figure 3.5: Chat Screen")
    
    add_custom_heading("3.6 Language Selection Screen", level=2)
    add_custom_paragraph("The Language Selection Screen allows users to set their default language choice (Step 6 of the onboarding sequence), configuring translation direction.")
    add_custom_image("language", "Figure 3.6: Language Selection Screen")

    # ------------------ CHAPTER 4 ------------------
    add_custom_heading("CHAPTER 4: DATABASE TABLES AND ER DIAGRAM", level=1)
    
    add_custom_heading("4.1 Users Table", level=2)
    add_custom_paragraph("The Users table contains profile details, preferred interface languages, and credentials.")
    add_custom_table(
        ["Field Name", "Type", "Description"],
        [
            ["user_id", "String", "Primary Key (Unique Email, Phone, or UID)"],
            ["username", "String", "User's Full Name / Display Name"],
            ["email", "String", "User's Registered Email Address"],
            ["preferred_language", "String", "Selected language code (e.g., 'en', 'hi', 'kn')"],
            ["created_at", "Timestamp", "Registration date and time record"]
        ]
    )
    
    add_custom_paragraph("Dummy Data Example:")
    add_custom_table(
        ["user_id", "username", "email", "preferred_language"],
        [
            ["U001", "Rahul", "rahul@gmail.com", "English (en)"],
            ["U002", "Priya", "priya@gmail.com", "Hindi (hi)"],
            ["U003", "Arjun", "arjun@gmail.com", "Kannada (kn)"]
        ]
    )
    
    add_custom_heading("4.2 Contacts Table", level=2)
    add_custom_paragraph("The Contacts table links users as conversational contacts, specifying contact names.")
    add_custom_table(
        ["contact_id", "owner_id", "contact_name"],
        [
            ["C001", "U001", "Priya (U002)"],
            ["C002", "U001", "Arjun (U003)"],
            ["C003", "U002", "Rahul (U001)"]
        ]
    )
    
    add_custom_heading("4.3 Messages Table", level=2)
    add_custom_paragraph("The Messages table logs individual chats. It holds foreign keys linking users, along with original text and translated text.")
    add_custom_table(
        ["message_id", "sender_id", "receiver_id", "original_text", "translated_text"],
        [
            ["M001", "U001", "U002", "Hello", "नमस्ते"],
            ["M002", "U002", "U001", "धन्यवाद", "Thank You"],
            ["M003", "U003", "U001", "ಹೇಗಿದ್ದೀಯ", "How are you"]
        ]
    )
    
    add_custom_heading("4.4 ER Diagram", level=2)
    add_custom_paragraph("The entity-relationship diagram below maps the schema structure, detailing the 1-to-many relationships from USERS to CONTACTS and MESSAGES tables.")
    add_custom_image("er_diagram", "Figure 4.1: Entity-Relationship Diagram")

    # ------------------ CHAPTER 5 ------------------
    add_custom_heading("CHAPTER 5: TESTING", level=1)
    add_custom_paragraph("Testing was performed across multiple scenarios to verify user login flows, real-time synchronization, and translation functionality.")
    
    add_custom_table(
        ["TC ID", "Test Scenario", "Input Data", "Expected Result", "Actual Result", "Status"],
        [
            ["TC_001", "Login with valid credentials", "Valid Email & Password", "Dashboard Opens", "Dashboard Opens", "PASS"],
            ["TC_002", "Invalid Password error handling", "Wrong Password input", "Error Message displayed", "Error Message displayed", "PASS"],
            ["TC_003", "Real-time Message Delivery", "Message Text entered", "Message instantly synchronized", "Message instantly synchronized", "PASS"],
            ["TC_004", "Automatic Language Translation", "English Message sent", "Hindi translation displayed", "Hindi translation displayed", "PASS"],
            ["TC_005", "User Logout sequence", "Click Logout option", "Login Screen opens", "Login Screen opens", "PASS"]
        ]
    )

    # ------------------ CHAPTER 6 ------------------
    add_custom_heading("CHAPTER 6: CONCLUSION AND FUTURE WORK", level=1)
    
    add_custom_heading("6.1 Conclusion", level=2)
    add_custom_paragraph(
        "The Smart Messenger application successfully demonstrates a functional approach to multilingual communication through real-time message translation. "
        "By integrating modern web structures (React.js + Vite), hybrid mobile compilation (Capacitor), and cloud backend services (Firebase Authentication and Cloud Firestore), "
        "the application offers a fast and robust environment that removes language boundaries. Dynamic speech-to-text dictation and text-to-speech feedback add usability, "
        "making the tool a comprehensive messaging utility."
    )
    
    add_custom_heading("6.2 Future Work", level=2)
    add_custom_paragraph("Future iterations of the application plan to incorporate the following features:")
    
    future_items = [
        "Voice Message Translation: Translating direct audio notes in addition to text translation.",
        "Video Calling Support: Standard face-to-face video chat capability with live captions.",
        "AI Chat Assistant Integration: Integrating local AI models to auto-respond or summarize long conversations.",
        "End-to-End Encryption: Enhancing message privacy with standard encryption protocols on Firestore data.",
        "Offline Translation Support: Implementing local dictionary models for translating without active internet connection.",
        "Group Chat Translation: Allowing multiple users with different preferred languages to text in a single group."
    ]
    for idx, item in enumerate(future_items, 1):
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.line_spacing = 1.15
        p_item.paragraph_format.space_after = Pt(3)
        run_num = p_item.add_run(f"{idx}. ")
        run_num.bold = True
        run_num.font.name = 'Times New Roman'
        run_text = p_item.add_run(item)
        run_text.font.name = 'Times New Roman'

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

    # ------------------ REFERENCES ------------------
    add_custom_heading("REFERENCES", level=1, space_before=18)
    
    refs = [
        "[1] React Documentation - https://react.dev/ - Component lifecycle and UI states.",
        "[2] Firebase Documentation - https://firebase.google.com/docs - Firebase Auth, Firestore real-time queries.",
        "[3] Capacitor Documentation - https://capacitorjs.com/docs - Hybrid mobile platform builds.",
        "[4] Google Translate API Documentation - Machine translation request structures.",
        "[5] Android Developer Documentation - https://developer.android.com/ - Build optimizations and APK compilation."
    ]
    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.line_spacing = 1.15
        p_ref.paragraph_format.space_after = Pt(4)
        run = p_ref.add_run(ref)
        run.font.name = 'Times New Roman'
        
    doc.save("Smart_Messenger_Project_Report.docx")
    print("Report generated successfully as 'Smart_Messenger_Project_Report.docx'.")

if __name__ == '__main__':
    create_report()
